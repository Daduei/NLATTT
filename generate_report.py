"""Generate Word report for Sections IV and V."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_color="D9E2F3"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_color)

    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)


def build_report():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Title
    title = doc.add_heading("Digital Signature Security Project", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Sections IV & V — Experimental Implementation, Results and Discussion")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph()

    # ===================== SECTION IV =====================
    doc.add_heading("IV. EXPERIMENTAL IMPLEMENTATION (Duy)", level=1)

    doc.add_heading("1. System Architecture", level=2)

    doc.add_heading("1.1 Overall Design", level=3)
    doc.add_paragraph(
        "The system implements a digital signature workflow for document integrity and authenticity. "
        "Instead of storing a raw binary signature, the implementation packages each signature inside "
        "a JSON metadata envelope (.sig file). This envelope records the signing algorithm, timestamp, "
        "signer identity, certificate fingerprint, document digest, and the Base64-encoded signature."
    )
    doc.add_paragraph("The high-level data flow is:")
    add_code_block(
        doc,
        "Test Document → SHA-256 Digest → Private Key Signing → JSON .sig Envelope\n"
        "                                                              ↓\n"
        "Public Key + X.509 Certificate → Verifier → Valid / Invalid Result"
    )
    p = doc.add_paragraph()
    p.add_run("Figure 1. ").bold = True
    p.add_run("System data flow from document signing to verification.")

    doc.add_paragraph("The project supports three signing algorithms:")
    add_table(doc,
        ["Algorithm", "Key Type", "Padding / Scheme", "Signature Size"],
        [
            ["RSA-PKCS1v15", "RSA-2048", "PKCS#1 v1.5 (RFC 8017)", "256 bytes"],
            ["RSA-PSS", "RSA-2048", "Probabilistic Signature Scheme", "256 bytes"],
            ["ECDSA-P-256", "EC P-256", "Elliptic Curve DSA (FIPS 186-5)", "71 bytes (DER)"],
    ])
    doc.add_paragraph(
        "RSA-PKCS1v15 and RSA-PSS share the same RSA-2048 key pair; only the padding scheme differs "
        "at signing time. ECDSA uses a separate P-256 key pair."
    )

    doc.add_heading("1.2 Module Structure", level=3)
    add_table(doc,
        ["Module", "Responsibility"],
        [
            ["main.py", "Orchestrates the full 7-step demonstration"],
            ["keygen.py", "Generates RSA-2048 and ECDSA P-256 key pairs"],
            ["certificate.py", "Creates self-signed X.509 certificates"],
            ["signer.py", "Signs documents and writes JSON .sig envelopes"],
            ["verifier.py", "Verifies signatures using digest + cryptographic checks"],
            ["algorithms.py", "Defines supported algorithms and RSA padding configuration"],
            ["utils.py", "Provides hex dumps, digest comparison, and result tables"],
            ["benchmark.py", "Measures signing/verification performance"],
            ["mitm_demo.py", "Simulates a Man-in-the-Middle document tampering attack"],
            ["test_scenario1–5.py", "Automated security test cases using pytest"],
        ],
    )
    p = doc.add_paragraph()
    p.add_run("Figure 2. ").bold = True
    p.add_run("Modular architecture of the experimental implementation.")

    doc.add_heading("1.3 Key Design Decisions", level=3)
    decisions = [
        "JSON envelope instead of raw bytes — Makes signatures self-describing and easier to inspect during experiments.",
        "Self-signed X.509 certificates — Binds each public key to a named identity (e.g., Dean of Faculty) without requiring an external Certificate Authority.",
        "Two-layer verification — Combines digest comparison (fast tamper indicator) with full cryptographic verification (public key check).",
        "Multi-algorithm support — Enables direct comparison of RSA and ECDSA in both security behavior and performance.",
    ]
    for d in decisions:
        doc.add_paragraph(d, style="List Number")

    # Section 2
    doc.add_heading("2. Implementation Environment and Tools", level=2)

    doc.add_heading("2.1 Environment", level=3)
    add_table(doc,
        ["Component", "Value"],
        [
            ["Operating System", "Windows 10/11"],
            ["Programming Language", "Python 3.14"],
            ["Cryptography Library", "cryptography ≥ 42.0.0"],
            ["Testing Framework", "pytest ≥ 8.0.0"],
            ["Hash Function", "SHA-256"],
            ["Key Formats", "PEM (PKCS#8 private, SubjectPublicKeyInfo public)"],
            ["Certificate Format", "X.509 (PEM-encoded, self-signed)"],
        ],
    )

    doc.add_heading("2.2 Generated Artifacts", level=3)
    add_table(doc,
        ["File", "Description"],
        [
            ["private_key_rsa.pem", "RSA-2048 private key"],
            ["public_key_rsa.pem", "RSA-2048 public key"],
            ["certificate_rsa.pem", "Self-signed X.509 certificate for RSA key"],
            ["private_key_ecdsa.pem", "ECDSA P-256 private key"],
            ["public_key_ecdsa.pem", "ECDSA P-256 public key"],
            ["certificate_ecdsa.pem", "Self-signed X.509 certificate for ECDSA key"],
            ["test_document.txt", "Original test document"],
            ["tampered_document.txt", "Modified document for tamper testing"],
            ["*.sig", "JSON signature envelope files"],
        ],
    )

    doc.add_heading("2.3 Test Document Content", level=3)
    doc.add_paragraph("The primary test document (test_document.txt) contains:")
    add_code_block(
        doc,
        "This is an official document issued by NEU.\n"
        "Amount: 1,000,000 VND\n"
        "Authorized by: Dean of Faculty"
    )
    p = doc.add_paragraph()
    p.add_run("SHA-256 digest of original document: ").bold = True
    p.add_run("7fae8634785588902a21f6d6030e8b6646750c66bc4882d196f62b1f38ff629c")

    # Section 3
    doc.add_heading("3. Step-by-Step Signing Implementation", level=2)

    doc.add_heading("Step 1: Key and Certificate Generation", level=3)
    doc.add_paragraph(
        "The system generates two key pairs: RSA-2048 (used by both RSA-PKCS1v15 and RSA-PSS) "
        "and ECDSA P-256. For each key pair, a self-signed X.509 certificate is created with:"
    )
    for item in [
        "Subject: CN=Dean of Faculty, OU=Faculty of IT, O=National Economics University, C=VN",
        "Validity: 365 days",
        "Key Usage: digital signature only",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    p = doc.add_paragraph()
    p.add_run("Figure 3. ").bold = True
    p.add_run("X.509 certificate binding a public key to the signer identity.")

    doc.add_heading("Step 2: Document Reading and Digest Computation", level=3)
    doc.add_paragraph(
        "The document is read as raw bytes. A SHA-256 digest is computed and displayed:"
    )
    add_code_block(doc, "[INFO] SHA-256 digest : 7fae8634785588902a21f6d6030e8b6646750c66bc4882d196f62b1f38ff629c")
    doc.add_paragraph(
        "This digest is stored inside the .sig envelope so the verifier can later detect tampering."
    )

    doc.add_heading("Step 3: Algorithm-Specific Signing", level=3)
    add_table(doc,
        ["Algorithm", "Signing Method"],
        [
            ["RSA-PKCS1v15", "private_key.sign(data, PKCS1v15(), SHA256())"],
            ["RSA-PSS", "private_key.sign(data, PSS(MGF1-SHA256), SHA256())"],
            ["ECDSA-P-256", "private_key.sign(data, ECDSA(SHA256()))"],
        ],
    )
    doc.add_paragraph(
        "The cryptography library hashes the document internally using SHA-256 before signing."
    )

    doc.add_heading("Step 4: JSON Metadata Envelope Creation", level=3)
    doc.add_paragraph("The signature is packaged into a .sig JSON file. Example structure:")
    add_code_block(
        doc,
        '{\n'
        '  "format_version": "1.0",\n'
        '  "algorithm": "RSA-PKCS1v15",\n'
        '  "hash_algorithm": "SHA-256",\n'
        '  "timestamp": "2026-06-11T20:27:51.090579+00:00",\n'
        '  "signer_id": "Dean of Faculty",\n'
        '  "certificate_fingerprint_sha256": "417cfbe5cfee80bb...",\n'
        '  "document_digest_sha256": "7fae8634785588902a21f6d6030e8b6646750c66bc4882d196f62b1f38ff629c",\n'
        '  "signature_b64": "w9ieBkxkaW5/vSIuRCkPi6a1EDCh56hxEbewIX/xbKYG...",\n'
        '  "signature_size_bytes": 256,\n'
        '  "key_type": "RSA-2048"\n'
        '}'
    )
    p = doc.add_paragraph()
    p.add_run("Table 1. ").bold = True
    p.add_run("Fields in the JSON signature envelope.")
    add_table(doc,
        ["Field", "Purpose"],
        [
            ["timestamp", "Records when the document was signed (UTC)"],
            ["algorithm", "Identifies the signing algorithm used"],
            ["signer_id", "Human-readable signer name from certificate CN"],
            ["certificate_fingerprint_sha256", "Links signature to a specific certificate"],
            ["document_digest_sha256", "SHA-256 hash at signing time"],
            ["signature_b64", "Base64-encoded cryptographic signature"],
        ],
    )

    doc.add_heading("Step 5: Output Inspection", level=3)
    doc.add_paragraph("After signing, the system prints:")
    for item in [
        "Document SHA-256 digest",
        "Hex dump of document bytes (first 64 bytes)",
        "Signature size in bytes",
        "First/last 16 bytes of the signature",
        "Hex dump preview of the signature",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    p = doc.add_paragraph()
    p.add_run("Figure 4. ").bold = True
    p.add_run("Console output showing document digest and signature inspection.")

    # Section 4
    doc.add_heading("4. Step-by-Step Verification Implementation", level=2)

    doc.add_heading("Step 1: Load Inputs", level=3)
    doc.add_paragraph("The verifier loads:")
    for item in [
        "The received document (bytes)",
        "The .sig JSON envelope",
        "The signer's public key (PEM)",
        "The X.509 certificate (optional, for identity display)",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Step 2: Digest Re-computation", level=3)
    doc.add_paragraph(
        "The verifier recomputes the SHA-256 digest of the received document and compares it "
        "with document_digest_sha256 stored in the envelope. If the digests differ, "
        "the document has been modified since signing."
    )

    doc.add_heading("Step 3: Cryptographic Verification", level=3)
    doc.add_paragraph(
        "The Base64 signature is decoded and verified against the document using the public key "
        "and the algorithm specified in the envelope. The library raises InvalidSignature if verification fails."
    )

    doc.add_heading("Step 4: Decision Logic", level=3)
    add_table(doc,
        ["Digest Match", "Crypto Verify", "Result"],
        [
            ["Yes", "Pass", "VALID — authentic and unmodified"],
            ["No", "Pass", "INVALID — digest mismatch proves tampering"],
            ["Any", "Fail", "INVALID — cryptographic verification failed"],
        ],
    )

    doc.add_heading("Step 5: Certificate and Identity Check", level=3)
    doc.add_paragraph(
        "If a certificate is provided, the verifier displays Subject, Issuer, Serial Number, "
        "validity period, signature algorithm, and SHA-256 fingerprint. It also checks whether "
        "signer_id in the envelope matches the certificate Common Name (CN)."
    )
    p = doc.add_paragraph()
    p.add_run("Figure 5. ").bold = True
    p.add_run("Two-layer verification process: digest check + cryptographic check.")

    # Section 5
    doc.add_heading("5. Test Documents and Security Testing Scenarios", level=2)

    doc.add_heading("5.1 Test Documents", level=3)
    add_table(doc,
        ["Document", "Purpose", "Content Change"],
        [
            ["test_document.txt", "Original signing test", "Amount: 1,000,000 VND"],
            ["tampered_document.txt", "Tamper detection test", "Amount: 9,000,000 VND"],
            ["report.bin", "Non-text file signing", "Simulated PDF header + binary data"],
        ],
    )

    doc.add_heading("5.2 Automated Security Tests (pytest)", level=3)
    doc.add_paragraph("Eight automated tests across five scenario files:")
    add_table(doc,
        ["Scenario", "File", "Security Property Tested", "Expected Result"],
        [
            ["1", "test_scenario1.py", "Valid signature on unmodified document", "PASS (valid)"],
            ["2", "test_scenario2.py", "Single-character tamper detection", "FAIL (invalid)"],
            ["3", "test_scenario3.py", "Wrong public key rejection", "FAIL (invalid)"],
            ["4", "test_scenario4.py", "Signature not reusable on different document", "FAIL (invalid)"],
            ["5a", "test_scenario5.py", "Binary file signing and verification", "PASS (valid)"],
            ["5b", "test_scenario5.py", "All 3 algorithms sign/verify roundtrip", "PASS (valid)"],
        ],
    )
    doc.add_paragraph("All 8 tests pass: 8 passed in 0.96s")

    doc.add_heading("5.3 Man-in-the-Middle Attack Simulation", level=3)
    doc.add_paragraph("The MITM demo (mitm_demo.py) models a realistic fraud scenario in three phases:")
    add_table(doc,
        ["Phase", "Actor", "Action"],
        [
            ["1", "Alice", "Signs an official transfer authorization (1,000,000 VND)"],
            ["2", "Attacker", "Intercepts document, changes amount to 9,000,000 VND, keeps original .sig"],
            ["3", "Bob", "Verifies received document — verification fails"],
        ],
    )
    doc.add_paragraph("Forensic evidence printed during the attack:")
    for item in [
        "SHA-256 digests before and after tampering",
        "Hex dumps of original vs tampered bytes",
        "Character-level diff showing exact modification positions",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5.4 Performance Benchmark Setup", level=3)
    doc.add_paragraph(
        "benchmark.py measures average signing and verification time over 100 iterations "
        "per algorithm using a fixed 279-byte payload."
    )

    doc.add_page_break()

    # ===================== SECTION V =====================
    doc.add_heading("V. RESULTS AND DISCUSSION (Duy – Dũng)", level=1)

    doc.add_heading("1. Signing Results", level=2)

    doc.add_heading("1.1 Successful Signing Across All Algorithms", level=3)
    add_table(doc,
        ["Algorithm", "Signature File", "Signature Size", "Key Type"],
        [
            ["RSA-PKCS1v15", "test_document.txt.rsa_pkcs1v15.sig", "256 bytes", "RSA-2048"],
            ["RSA-PSS", "test_document.txt.rsa_pss.sig", "256 bytes", "RSA-2048"],
            ["ECDSA-P-256", "test_document.txt.ecdsa_p256.sig", "71 bytes", "EC P-256"],
        ],
    )
    p = doc.add_paragraph()
    p.add_run("Table 2. ").bold = True
    p.add_run("Signing results for three algorithms on the same document.")
    doc.add_paragraph(
        "ECDSA produces the smallest signature (71 bytes vs 256 bytes for RSA), "
        "which is an advantage for bandwidth-constrained systems."
    )

    doc.add_heading("1.2 SHA-256 Digest Consistency", level=3)
    doc.add_paragraph(
        "All three algorithms produce the same document digest because they hash the identical "
        "document with SHA-256 before signing:"
    )
    add_code_block(doc, "7fae8634785588902a21f6d6030e8b6646750c66bc4882d196f62b1f38ff629c")
    doc.add_paragraph("Only the cryptographic signature bytes differ between algorithms.")

    doc.add_heading("1.3 JSON Envelope Completeness", level=3)
    doc.add_paragraph(
        "Each .sig file contains complete metadata: timestamp, signer ID (Dean of Faculty), "
        "certificate fingerprint, and document digest. This makes each signature self-contained "
        "and auditable without external lookup tables."
    )
    p = doc.add_paragraph()
    p.add_run("Figure 6. ").bold = True
    p.add_run("Sample JSON .sig envelope for RSA-PKCS1v15.")

    doc.add_heading("2. Verification Results", level=2)

    doc.add_heading("2.1 Original Document Verification", level=3)
    add_table(doc,
        ["Algorithm", "Digest Match", "Crypto Verify", "Result"],
        [
            ["RSA-PKCS1v15", "Yes", "Pass", "VALID"],
            ["RSA-PSS", "Yes", "Pass", "VALID"],
            ["ECDSA-P-256", "Yes", "Pass", "VALID"],
        ],
    )
    add_code_block(doc, "[RESULT] Signature VALID — Document is authentic and unmodified.")
    doc.add_paragraph("All three algorithms correctly confirmed document authenticity.")

    doc.add_heading("2.2 Certificate Identity Binding", level=3)
    add_code_block(
        doc,
        "Subject    : CN=Dean of Faculty, OU=Faculty of IT, O=National Economics University, C=VN\n"
        "Valid From : 2026-06-11T20:27:51+00:00\n"
        "Valid To   : 2027-06-11T20:27:51+00:00\n"
        "Signature  : sha256WithRSAEncryption"
    )
    doc.add_paragraph(
        "The signer_id field in the .sig envelope matches the certificate CN, confirming identity consistency."
    )
    p = doc.add_paragraph()
    p.add_run("Figure 7. ").bold = True
    p.add_run("X.509 certificate details displayed during verification.")

    doc.add_heading("3. Tampered Document Verification Results", level=2)

    doc.add_heading("3.1 Tamper Detection", level=3)
    doc.add_paragraph("After changing the amount from 1,000,000 to 9,000,000 VND, verification failed for all three algorithms:")
    add_table(doc,
        ["Algorithm", "Stored Digest (prefix)", "Current Digest (prefix)", "Result"],
        [
            ["RSA-PKCS1v15", "7fae8634...", "26a5c91b...", "INVALID"],
            ["RSA-PSS", "7fae8634...", "26a5c91b...", "INVALID"],
            ["ECDSA-P-256", "7fae8634...", "26a5c91b...", "INVALID"],
        ],
    )
    p = doc.add_paragraph()
    p.add_run("Original digest: ").bold = True
    p.add_run("7fae8634785588902a21f6d6030e8b6646750c66bc4882d196f62b1f38ff629c")
    p = doc.add_paragraph()
    p.add_run("Tampered digest: ").bold = True
    p.add_run("26a5c91b2189559ffdaa9cacb2317102a365960a4e58a0525a90f5a66d54a2ae")
    doc.add_paragraph(
        "60 out of 64 hex characters differ — demonstrating the SHA-256 avalanche effect, "
        "where even a small text change produces a completely different hash."
    )
    p = doc.add_paragraph()
    p.add_run("Figure 8. ").bold = True
    p.add_run("SHA-256 digest comparison before and after tampering.")

    doc.add_heading("3.2 MITM Attack Simulation Result", level=3)
    add_table(doc,
        ["Phase", "Outcome"],
        [
            ["Alice signs original (1M VND)", "Signature created successfully"],
            ["Attacker changes amount to 9M VND", "Digest changes completely"],
            ["Bob verifies tampered document", "Attack BLOCKED — verification fails"],
        ],
    )
    add_code_block(
        doc,
        "[OK] Attack BLOCKED — Digital signature detected the tampering.\n"
        "     The SHA-256 digest changed, so the signature no longer matches.\n"
        "     Bob knows the document was modified in transit."
    )

    doc.add_heading("3.3 Other Attack Scenarios", level=3)
    add_table(doc,
        ["Attack Type", "Test", "Result"],
        [
            ["Wrong public key", "test_scenario3.py", "Rejected — INVALID"],
            ["Signature reuse on different document", "test_scenario4.py", "Rejected — INVALID"],
            ["Single-character modification", "test_scenario2.py", "Rejected — INVALID"],
        ],
    )
    p = doc.add_paragraph()
    p.add_run("Table 3. ").bold = True
    p.add_run("Security test scenario results.")

    doc.add_heading("4. Performance Evaluation Results", level=2)

    doc.add_heading("4.1 Benchmark Methodology", level=3)
    for item in [
        "Iterations: 100 per operation (sign and verify)",
        "Payload: 279 bytes (fixed random content)",
        "Timing: time.perf_counter() — average milliseconds per operation",
        "Environment: Local Windows machine, Python 3.14",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4.2 Performance Comparison Table", level=3)
    p = doc.add_paragraph()
    p.add_run("Table 4. ").bold = True
    p.add_run("Performance benchmark results (100 iterations, average per operation).")
    add_table(doc,
        ["Algorithm", "Sign (ms)", "Verify (ms)", "Sig Size (B)", "Key (bits)"],
        [
            ["RSA-PKCS1v15", "0.288", "0.028", "256", "2048"],
            ["RSA-PSS", "0.274", "0.029", "256", "2048"],
            ["ECDSA-P-256", "0.027", "0.065", "71", "256"],
        ],
    )
    p = doc.add_paragraph()
    p.add_run("Figure 9. ").bold = True
    p.add_run("Performance comparison chart of signing and verification times.")

    doc.add_heading("4.3 Analysis", level=3)
    add_table(doc,
        ["Metric", "Winner", "Value", "Explanation"],
        [
            ["Fastest signing", "ECDSA-P-256", "0.027 ms", "Elliptic curve ops faster than RSA exponentiation"],
            ["Fastest verification", "RSA-PKCS1v15", "0.028 ms", "RSA public-key ops faster than EC verify"],
            ["Smallest signature", "ECDSA-P-256", "71 bytes", "EC signatures ~3.6× smaller than RSA-2048"],
            ["Smallest key", "ECDSA-P-256", "256 bits", "P-256 provides ~128-bit security"],
        ],
    )
    doc.add_paragraph(
        "Discussion: ECDSA-P-256 is the best choice when signing speed and signature size matter "
        "(e.g., IoT, mobile). RSA-PKCS1v15 is preferable when verification speed is critical. "
        "RSA-PSS is recommended by modern standards (PKCS#1 v2.1) for new deployments due to "
        "its probabilistic padding, though performance is similar to PKCS1v15."
    )

    doc.add_heading("5. Workflow Analysis and Security Evaluation", level=2)

    doc.add_heading("5.1 End-to-End Workflow", level=3)
    add_table(doc,
        ["Step", "Action", "Outcome"],
        [
            ["1", "Generate RSA + ECDSA keys and X.509 certificates", "Keys and certs created"],
            ["2", "Create test document", "test_document.txt written"],
            ["3", "Sign with all 3 algorithms", "3 .sig files created"],
            ["4", "Verify original document", "All 3 PASS"],
            ["5", "Tamper document and verify", "All 3 FAIL"],
            ["6", "Performance benchmark", "Table printed"],
            ["7", "MITM attack simulation", "Attack blocked"],
        ],
    )
    p = doc.add_paragraph()
    p.add_run("Figure 10. ").bold = True
    p.add_run("Complete 7-step demonstration workflow.")

    doc.add_heading("5.2 Security Properties Demonstrated", level=3)
    add_table(doc,
        ["Property", "Mechanism", "Evidence"],
        [
            ["Integrity", "SHA-256 digest + signature binding", "Tampered doc fails verification"],
            ["Authenticity", "Private key signing + X.509 identity", "Only holder of private key can sign"],
            ["Non-repudiation", "Signature tied to signer ID + timestamp", ".sig envelope records who signed and when"],
            ["Unforgeability", "Wrong key / reused signature rejected", "Scenarios 3 and 4 fail correctly"],
        ],
    )

    doc.add_heading("5.3 Attack Resistance Summary", level=3)
    doc.add_paragraph("The implementation successfully defended against:")
    attacks = [
        "Content tampering — changing even one digit invalidates the signature",
        "Key substitution — verifying with a different public key fails",
        "Signature transplant — using Document A's signature on Document B fails",
        "Man-in-the-Middle modification — intercepted and modified documents are detected at verification time",
    ]
    for a in attacks:
        doc.add_paragraph(a, style="List Number")

    doc.add_heading("6. Limitations", level=2)
    add_table(doc,
        ["#", "Limitation", "Impact"],
        [
            ["1", "Self-signed certificates only", "No real CA chain; trust is not externally verifiable"],
            ["2", "Local simulation only", "MITM attack is simulated in software, not over a real network"],
            ["3", "Unencrypted private keys", "PEM files stored without passphrase protection"],
            ["4", "Limited algorithm set", "Only RSA-2048 and ECDSA P-256; no Ed25519 or post-quantum"],
            ["5", "Environment-dependent benchmarks", "Performance numbers vary by CPU and system load"],
            ["6", "Functional tests, not formal audit", "pytest scenarios demonstrate correctness, not penetration testing"],
            ["7", "No timestamp authority", "Timestamp is local system clock, not from a trusted TSA"],
            ["8", "No certificate revocation", "No CRL or OCSP mechanism to invalidate compromised certificates"],
        ],
    )
    doc.add_paragraph(
        "Despite these limitations, the implementation is sufficient for demonstrating core digital "
        "signature principles, comparing algorithms, and evaluating security properties in an academic setting."
    )

    doc.add_page_break()

    # Figures appendix
    doc.add_heading("Appendix: Suggested Figures for the Report", level=1)
    add_table(doc,
        ["Figure", "Caption", "Source"],
        [
            ["Figure 1", "System data flow diagram", "Section IV.1 — draw from data flow description"],
            ["Figure 2", "Module architecture diagram", "Section IV.1.2 — draw from module table"],
            ["Figure 3", "X.509 certificate output screenshot", "Run: py keygen.py"],
            ["Figure 4", "Signing output with hex dump", "Run: py signer.py"],
            ["Figure 5", "Two-layer verification flowchart", "Section IV.4 — draw digest + crypto check flow"],
            ["Figure 6", "JSON .sig envelope sample", "test_document.txt.rsa_pkcs1v15.sig"],
            ["Figure 7", "Certificate info during verification", "Run: py verifier.py"],
            ["Figure 8", "Digest comparison after tampering", "Run: py main.py Step 5"],
            ["Figure 9", "Performance bar chart", "Plot Table 4 data in Excel"],
            ["Figure 10", "7-step workflow diagram", "Section V.5.1 — draw from workflow table"],
        ],
    )

    output_path = "Report_Section_IV_V.docx"
    doc.save(output_path)
    print(f"[OK] Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_report()
